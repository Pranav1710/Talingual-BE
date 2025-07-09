# services/exporter/docx_exporter.py

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from bs4 import BeautifulSoup, NavigableString
from io import BytesIO
import base64
import tempfile


def convert_html_to_docx(html: str, config: dict) -> str:
    doc = Document()

    # Config values
    font = config.get("fontFamily", "Arial")
    size = int(config.get("fontSize", "11px").replace("px", ""))
    spacing = float(config.get("lineSpacing", 1.15))
    logo_width = int(config.get("logoSize", "250px").replace("px", ""))
    show_logo = config.get("showLogo", True)
    # Style setup
    style = doc.styles["Normal"]
    style.font.name = font
    style.font.size = Pt(size)

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    soup = BeautifulSoup(html, "html.parser")

    # Inject logo if present
    if show_logo:
        insert_logo_from_html(doc, soup, logo_width)

    last_tag = None
    for elem in soup.find_all(["h2", "p", "ul", "li"], recursive=True):
        if isinstance(elem, NavigableString):
            continue

        tag = elem.name
        class_list = elem.get("class", [])
        text = elem.get_text(strip=True)
        if not text:
            continue

        if tag == "h2":
            para = doc.add_paragraph()
            run = para.add_run(text.upper())
            run.bold = True
            run.font.size = Pt(size + 1)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.space_before = Pt(12)
            last_tag = "h2"

        elif tag == "p":
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.size = Pt(size)
            para.paragraph_format.line_spacing = spacing
            apply_paragraph_spacing(para, class_list)
            last_tag = "p"

        elif tag == "li":
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(text)
            run.font.size = Pt(size)
            para.paragraph_format.line_spacing = spacing
            para.paragraph_format.space_after = Pt(5)
            last_tag = "li"

    # Add final blank line
    doc.add_paragraph()

    # Save to temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)
    return temp_file.name


def insert_logo_from_html(doc, soup, width_px: int):
    """
    Looks for <img class='resume-logo'> in the HTML and injects it into the DOCX.
    """
    logo_tag = soup.find("img", class_="resume-logo")
    if logo_tag and "base64" in logo_tag.get("src", ""):
        base64_data = logo_tag["src"].split("base64,")[1]
        logo_bytes = BytesIO(base64.b64decode(base64_data))
        doc.add_picture(logo_bytes, width=Inches(width_px / 96))
        doc.add_paragraph()


def apply_paragraph_spacing(para, class_list):
    """
    Applies spacing rules to a paragraph based on its HTML class.
    """
    spacing_map = {
        "profile-paragraph": (12, 8),
        "profile-paragraph-note": (0, 14),
        "work-dates": (12, 2),
        "work-title": (0, 8),
        "edu-line-1": (12, 0),
        "edu-line-2": (0, 4),
        "info-field": (0, 3),
    }

    for cls in class_list:
        if cls in spacing_map:
            before, after = spacing_map[cls]
            para.paragraph_format.space_before = Pt(before)
            para.paragraph_format.space_after = Pt(after)
            return

    # Default spacing
    para.paragraph_format.space_after = Pt(5)
    
from flask import request, session, send_file
from utils.html_parser import extract_filename_from_html

def generate_docx():
    """
    Flask route handler that converts styled HTML to .docx and returns it.
    """
    html = request.json.get("html", "")
    config = session.get("resume_default_config", {})

    path = convert_html_to_docx(html, config)
    filename = extract_filename_from_html(html).replace(".pdf", ".docx")

    return send_file(path, as_attachment=True, download_name=filename)
